"""
Document processing service with OpenTelemetry metrics
"""
from pathlib import Path
import logging
from monitoring.tracing import create_span
from monitoring.metrics import (
    document_processing_duration,
    document_chars_processed
)
import time

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for processing documents"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}
    MAX_SIZE_MB = 20
    
    @staticmethod
    def validate_document(file_path: str) -> dict:
        """Validate uploaded document"""
        path = Path(file_path)
        
        # Check extension
        if path.suffix.lower() not in DocumentService.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Unsupported format: {path.suffix}. "
                f"Allowed: {DocumentService.ALLOWED_EXTENSIONS}"
            )
        
        # Check file size
        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > DocumentService.MAX_SIZE_MB:
            raise ValueError(f"File too large: {size_mb:.2f}MB. Max: {DocumentService.MAX_SIZE_MB}MB")
        
        return {
            "extension": path.suffix,
            "size_mb": round(size_mb, 2),
            "name": path.name
        }
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Parse PDF file"""
        start_time = time.time()
        
        with create_span("parse_pdf", {"file_path": file_path}):
            try:
                import PyPDF2
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page.extract_text() + "\n"
                    
                    # Record metrics
                    duration = time.time() - start_time
                    document_processing_duration.labels(document_type="pdf").observe(duration)
                    document_chars_processed.inc(len(text))
                    
                    logger.info(f"✅ Parsed PDF: {len(pdf_reader.pages)} pages, {len(text)} chars in {duration:.2f}s")
                    return text.strip()
                    
            except Exception as e:
                logger.error(f"❌ PDF parsing failed: {e}")
                raise ValueError(f"Could not parse PDF: {e}")
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Parse DOCX file"""
        start_time = time.time()
        
        with create_span("parse_docx", {"file_path": file_path}):
            try:
                import docx
                
                doc = docx.Document(file_path)
                text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                
                # Record metrics
                duration = time.time() - start_time
                document_processing_duration.labels(document_type="docx").observe(duration)
                document_chars_processed.inc(len(text))
                
                logger.info(f"✅ Parsed DOCX: {len(doc.paragraphs)} paragraphs, {len(text)} chars in {duration:.2f}s")
                return text.strip()
                
            except Exception as e:
                logger.error(f"❌ DOCX parsing failed: {e}")
                raise ValueError(f"Could not parse DOCX: {e}")
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Parse TXT/MD file"""
        start_time = time.time()
        
        with create_span("parse_txt", {"file_path": file_path}):
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                
                # Record metrics
                duration = time.time() - start_time
                document_processing_duration.labels(document_type="txt").observe(duration)
                document_chars_processed.inc(len(text))
                
                logger.info(f"✅ Parsed TXT: {len(text)} chars in {duration:.2f}s")
                return text.strip()
                
            except UnicodeDecodeError:
                # Try with different encoding
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                duration = time.time() - start_time
                document_processing_duration.labels(document_type="txt").observe(duration)
                document_chars_processed.inc(len(text))
                
                logger.info(f"✅ Parsed TXT (latin-1): {len(text)} chars in {duration:.2f}s")
                return text.strip()
                
            except Exception as e:
                logger.error(f"❌ TXT parsing failed: {e}")
                raise ValueError(f"Could not parse TXT: {e}")
    
    @staticmethod
    def parse_document(file_path: str) -> str:
        """Auto-detect and parse document"""
        with create_span("parse_document", {"file_path": file_path}):
            path = Path(file_path)
            extension = path.suffix.lower()
            
            if extension == '.pdf':
                return DocumentService.parse_pdf(file_path)
            elif extension == '.docx':
                return DocumentService.parse_docx(file_path)
            elif extension in ['.txt', '.md']:
                return DocumentService.parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")


# Singleton instance
document_service = DocumentService()