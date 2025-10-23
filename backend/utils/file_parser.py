"""
Utility for parsing documents (PDF, DOCX, TXT, MD)
"""
from pathlib import Path
import logging
from typing import Optional

import PyPDF2
import docx

logger = logging.getLogger(__name__)


class FileParser:
    """Utility class to parse different file types"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    @staticmethod
    def parse(file_path: str) -> Optional[str]:
        """
        Parse a document file and return extracted text.
        Auto-detects file type from extension.

        Args:
            file_path (str): Path to the file

        Returns:
            str: Extracted text content or None
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in FileParser.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        if ext == ".pdf":
            return FileParser._parse_pdf(file_path)
        elif ext == ".docx":
            return FileParser._parse_docx(file_path)
        elif ext in [".txt", ".md"]:
            return FileParser._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    # ---- Internal parsers ----

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Extract text from PDF using PyPDF2"""
        try:
            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

            logger.info(f"✅ Parsed PDF ({len(reader.pages)} pages, {len(text)} chars)")
            return text.strip()

        except Exception as e:
            logger.error(f"❌ Failed to parse PDF {file_path}: {e}")
            raise

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            doc = docx.Document(file_path)
            text = "\n\n".join(
                [p.text for p in doc.paragraphs if p.text.strip()]
            )
            logger.info(f"✅ Parsed DOCX ({len(doc.paragraphs)} paragraphs, {len(text)} chars)")
            return text.strip()

        except Exception as e:
            logger.error(f"❌ Failed to parse DOCX {file_path}: {e}")
            raise

    @staticmethod
    def _parse_text(file_path: str) -> str:
        """Read text or Markdown files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            logger.info(f"✅ Parsed TXT ({len(text)} chars)")
            return text.strip()

        except Exception as e:
            logger.error(f"❌ Failed to parse TXT {file_path}: {e}")
            raise


# Singleton instance
file_parser = FileParser()
